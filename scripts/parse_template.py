"""
parse_template.py

解析广发研究报告Word模板
输出 report_style.json

依赖:
pip install python-docx
"""


from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path
import json


# ==============================
# 路径配置
# ==============================
ROOT_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = Path(
    ROOT_DIR

    / "template"

    / "总量通用.docx"
)

OUTPUT_PATH = Path(
    ROOT_DIR

    / "config"

    / "report_style.json"
)


# ==============================
# 工具函数
# ==============================

def font_to_dict(font):
    """
    提取字体信息
    """

    result = {}

    if font.name:
        result["name"] = font.name

    if font.size:
        result["size_pt"] = round(
            font.size.pt,
            2
        )

    if font.bold is not None:
        result["bold"] = font.bold

    if font.italic is not None:
        result["italic"] = font.italic


    if font.color and font.color.rgb:
        result["color"] = str(
            font.color.rgb
        )


    return result



def paragraph_format_to_dict(fmt):
    """
    提取段落格式
    """

    result = {}

    if fmt.space_before:
        result["space_before"] = (
            fmt.space_before.pt
        )

    if fmt.space_after:
        result["space_after"] = (
            fmt.space_after.pt
        )

    if fmt.line_spacing:
        if isinstance(
            fmt.line_spacing,
            float
        ):
            result["line_spacing"] = (
                fmt.line_spacing
            )
        else:
            try:
                result["line_spacing_pt"] = (
                    fmt.line_spacing.pt
                )
            except:
                pass


    if fmt.first_line_indent:
        result["first_line_indent"] = (
            fmt.first_line_indent.pt
        )


    if fmt.left_indent:
        result["left_indent"] = (
            fmt.left_indent.pt
        )


    return result



# ==============================
# 解析页面
# ==============================

def parse_page(doc):

    section = doc.sections[0]

    return {

        "page_width_cm":
            round(
                section.page_width.cm,
                2
            ),

        "page_height_cm":
            round(
                section.page_height.cm,
                2
            ),

        "margin":{

            "top_cm":
                round(
                    section.top_margin.cm,
                    2
                ),

            "bottom_cm":
                round(
                    section.bottom_margin.cm,
                    2
                ),

            "left_cm":
                round(
                    section.left_margin.cm,
                    2
                ),

            "right_cm":
                round(
                    section.right_margin.cm,
                    2
                )
        }
    }



# ==============================
# 解析 styles
# ==============================

def parse_styles(doc):

    styles = {}

    for style in doc.styles:


        item = {

            "type":
                str(style.type),

            "font":
                font_to_dict(
                    style.font
                ),

        }


        try:
            item["paragraph"] = (
                paragraph_format_to_dict(
                    style.paragraph_format
                )
            )

        except:
            pass


        styles[
            style.name
        ] = item


    return styles



# ==============================
# 解析文档统计信息
# ==============================

def parse_structure(doc):

    stats = {

        "paragraph_count":0,

        "table_count":0,

        "image_count":0
    }


    for p in doc.paragraphs:

        stats["paragraph_count"] += 1


    stats["table_count"] = len(
        doc.tables
    )


    # 图片数量
    rels = doc.part.rels

    for rel in rels.values():

        if "image" in rel.target_ref:
            stats["image_count"] += 1


    return stats



# ==============================
# 主函数
# ==============================

def main():

    print("Loading template...")

    doc = Document(
        TEMPLATE_PATH
    )


    result = {


        "template":
            TEMPLATE_PATH.name,


        "page":
            parse_page(doc),


        "styles":
            parse_styles(doc),


        "statistics":
            parse_structure(doc)

    }


    OUTPUT_PATH.parent.mkdir(
        exist_ok=True
    )


    with open(
        OUTPUT_PATH,
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            result,
            f,
            ensure_ascii=False,
            indent=4
        )


    print(
        "Done!"
    )

    print(
        f"Saved to {OUTPUT_PATH}"
    )



if __name__ == "__main__":

    main()